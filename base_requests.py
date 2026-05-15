"""
base_requests.py
Бизнес-логика: обновление фильмов, работа с платежами, регистрация пользователей.

Все вызовы бота заменены на sync-bridge из config.py:
  await notify_admin(text)           — отправить сообщение администратору
  send_message_sync(id, text)  — отправить сообщение пользователю
  send_document_sync(...)      — отправить документ (PDF-билет)
"""
import asyncio
import datetime
import hashlib
import json
import os
import re
import subprocess
import time
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeout
from io import BytesIO

import psycopg2
import requests
from docx import Document
from loguru import logger
from maxapi.types import ButtonsPayload, LinkButton
from yookassa import Payment, Configuration

from config import (
    db_path,
    info_log_file,
    kinopoisk_token,
    notify_admin,
    pochta_bank_token,
    root_path,
    send_document_sync,
    send_message_sync,
    url,
    url_kino_baza,
    url_prokultura,
    youkassa_shop_id,
    youkassa_secret_key,
    validation_url, url_server, get_loop,
)
from pkg.log import CustomLogger

CustomLogger().add_logger(info_log_file, __name__)

ADMIN_ID = 8099031
ADMIN_ID_CHAT = 284784629
ADMIN_ID_2 = 1013689498  # Антон


# ─── Главный цикл обновления ──────────────────────────────────────────────────
def film_update_main(loop):
    i = 1
    kinopoisk_enabled = True
    last_check_date = datetime.date.today()

    while True:
        start = time.time()

        try:
            # reset daily flag
            today = datetime.date.today()
            if today != last_check_date:
                kinopoisk_enabled = True
                last_check_date = today
                logger.info("Новый день — kinopoisk_enabled=True")

            kinopoisk_enabled = run_tasks(i, kinopoisk_enabled, loop)

        except Exception as e:
            logger.exception("Ошибка в film_update_main")
            asyncio.run_coroutine_threadsafe(notify_admin(f"Ошибка в film_update_main\n{e}"), loop)

        finally:
            i = 1 if i > 10000 else i + 1

            # фиксированный тик, а не просто sleep(5)
            elapsed = time.time() - start
            sleep_time = max(0, 5 - elapsed)
            time.sleep(sleep_time)


def run_tasks(i, kinopoisk_enabled, loop):
    # safe_execute(get_show_info, "get_show_info", loop=loop)

    if i % 2 == 0 or i == 1:
        safe_execute(all_show_request, "all_show_request", loop=loop)

    if i % 4 == 0 or i == 1:
        safe_execute(all_performances_request, "all_performances_request", loop=loop)

    if i % 10 == 0 or i == 1:
        safe_execute(what_show_can_be_sell_pushkin_card, "what_show_can_be_sell_pushkin_card", loop=loop)

    if kinopoisk_enabled:
        res = safe_execute(get_kinopoisk_info, "get_kinopoisk_info", loop=loop)
        if res is False:
            kinopoisk_enabled = False

    safe_execute(unblock_5_min, "unblock_5_min", loop=loop)

    return kinopoisk_enabled


def safe_execute(func, name, **kwargs):
    try:
        # logger.info(f"{name} start")
        return func(**kwargs)
    except Exception as e:
        logger.exception(f"{name} error: {e}")
        return None
    # finally:
    #     logger.info(f"{name} stop")


def process_orders(loop):
    while True:
        try:
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute("SELECT payment_id FROM orders WHERE status = 3;")
                    orders = curs.fetchall()

            for order in orders:
                asyncio.run_coroutine_threadsafe(check_payment_status(order[0]), loop)

        except Exception as e:
            logger.exception("Ошибка при получении заказов")
            asyncio.run_coroutine_threadsafe(notify_admin(f'Ошибка при получении заказов\n{e}'), loop)

        time.sleep(10)


# ─── Обновление данных о фильмах ──────────────────────────────────────────────
def all_show_request(loop):
    params = {"sp": "Wga_GetShow", "df": "J"}
    response = requests.get(url_kino_baza, params=params)

    if response.status_code == 200:
        try:
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    for show in response.json():
                        try:
                            show['ShowName'] = (
                                show['ShowName']
                                .replace('&quot;', '')
                                .replace('«', '')
                                .replace('»', '')
                            )
                            curs.execute(
                                """INSERT INTO show (show_id, name, duration)
                                   VALUES (%s, %s, %s)
                                   ON CONFLICT (show_id) DO NOTHING;""",
                                (show['IdShow'], show['ShowName'], show['Duration']),
                            )
                            curs.execute(
                                "UPDATE show SET name = %s, duration = %s WHERE show_id = %s",
                                (show['ShowName'], show['Duration'], show['IdShow']),
                            )
                        except Exception as e:
                            logger.exception(f"Ошибка при вставке шоу {show.get('IdShow')}: {e}")
        except json.JSONDecodeError as e:
            asyncio.run_coroutine_threadsafe(notify_admin(f'Ошибка декодирования JSON в all_show_request: {e}'), loop)
        except Exception as e:
            asyncio.run_coroutine_threadsafe(notify_admin(f'Ошибка базы по запросу всех show\n{e}'), loop)
    else:
        asyncio.run_coroutine_threadsafe(notify_admin(f'Ошибка all_show_request: код {response.status_code}'), loop)


def get_show_info(loop):
    try:
        with psycopg2.connect(db_path) as conn:
            with conn.cursor() as curs:
                curs.execute("SELECT show_id, name FROM show WHERE kinopoisk_id IS NULL;")
                shows = curs.fetchall()
                for show in shows:
                    params = {"sp": "Wga_GetShowInfo", "idShow": show[0], "df": "J"}
                    response = requests.get(url_kino_baza, params=params)
                    if response.status_code == 200:
                        data = response.json()
                        kinopoisk_id = data['Remark']
                        if kinopoisk_id == " ":
                            kinopoisk_id = None
                        curs.execute(
                            "UPDATE show SET kinopoisk_id = %s WHERE show_id = %s",
                            (kinopoisk_id, show[0]),
                        )
                        conn.commit()
    except Exception as e:
        logger.exception("Ошибка get_show_info")
        asyncio.run_coroutine_threadsafe(notify_admin(f'Ошибка Wga_GetShowInfo\n{e}'), loop)


def normalize_name(name: str) -> str:
    name = name.lower()
    name = re.sub(r'[^\w\dа-яё]', '', name, flags=re.UNICODE)
    return name


def what_show_can_be_sell_pushkin_card(loop):
    try:
        params = {
            "organizations": "31698",
            'limit': '100',
            'categories': 'kino',
            "isPushkinsCard": "true",
            "apiKey": "txjo5hdvmmgqs7frjb64",
            'status': 'accepted',
        }
        response = requests.get(url_prokultura, params=params)

        if response.status_code == 200:
            data = response.json()
            with open('data_prokult.json', 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=4, ensure_ascii=False)

            accepted_films_list = []
            pu_number_list = {}
            id_procult_list = {}

            for event in data['events']:
                match = re.search(r'«(.+)', event['name'])
                try:
                    film_name = match.group(1) if match else event['name']
                    film_name = normalize_name(film_name)
                except Exception as e:
                    logger.exception(f"Ошибка в what_show_can_be_sell_pushkin_card: {e}")
                    continue
                pu_number_list[film_name] = event['rentalCertificate'][0]
                id_procult_list[film_name] = event['_id']
                accepted_films_list.append(film_name)

            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute("SELECT DISTINCT name FROM show;")
                    shows = curs.fetchall()
                    for show in shows:
                        film_name = normalize_name(show[0])
                        if film_name in accepted_films_list:
                            curs.execute(
                                "UPDATE show SET pushkin_card = 1, pu_number = %s, id_procult = %s WHERE name = %s",
                                (pu_number_list[film_name], id_procult_list[film_name], show[0]),
                            )
                        else:
                            curs.execute(
                                "UPDATE show SET pushkin_card = 0 WHERE name = %s", (show[0],)
                            )
        else:
            asyncio.run_coroutine_threadsafe(
                notify_admin(f'Ошибка what_show_can_be_sell_pushkin_card: код {response.status_code} '
                             f'текст {response.text}'), loop)
    except Exception as e:
        logger.exception("Ошибка what_show_can_be_sell_pushkin_card")
        asyncio.run_coroutine_threadsafe(
            notify_admin(f'Ошибка what_show_can_be_sell_pushkin_card\n{e}'), loop)


def get_kinopoisk_info(loop):
    try:
        with psycopg2.connect(db_path) as conn:
            with conn.cursor() as curs:
                curs.execute(
                    "SELECT kinopoisk_id FROM show "
                    "WHERE kinopoisk_id IS NOT NULL AND (poster IS NULL OR poster = '');"
                )
                shows = curs.fetchall()
                for show in shows:
                    kinopoisk_id = show[0]
                    response = requests.get(
                        f'https://api.kinopoisk.dev/v1/movie/{kinopoisk_id}',
                        params={"token": kinopoisk_token},
                    )
                    if response.status_code == 403:
                        try:
                            err = response.json()
                        except Exception:
                            err = {}
                        if "суточный лимит" in err.get("message", ""):
                            logger.error("Суточный лимит Кинопоиск исчерпан")
                            asyncio.run_coroutine_threadsafe(
                                notify_admin("Суточный лимит Кинопоиск исчерпан. Обновление остановлено."), loop)
                            return False
                        continue
                    if response.status_code == 200:
                        data = response.json()
                        poster = (data.get("poster") or {}).get("url", "")
                        curs.execute(
                            "UPDATE show SET description = %s, poster = %s, kp_rating = %s "
                            "WHERE kinopoisk_id = %s",
                            (data.get('description'), poster, (data.get('rating') or {}).get('kp'), kinopoisk_id),
                        )
                        conn.commit()
        return True
    except Exception as e:
        logger.exception("Ошибка get_kinopoisk_info")
        asyncio.run_coroutine_threadsafe(
            notify_admin(f'Ошибка get_kinopoisk_info\n{e}'), loop)
        return True


def all_performances_request(loop):
    params = {"sp": "Wga_GetPerformance", "df": "J"}
    response = requests.get(url_kino_baza, params=params)

    if response.status_code == 200:
        try:
            data = response.json()
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute("DELETE FROM performance")
                    for item in data:
                        try:
                            try:
                                dt = datetime.datetime.strptime(item['DateTime'], "%B %d %Y %I:%M:%S:%f%p")
                            except ValueError:
                                dt = datetime.datetime.strptime(item['DateTime'], "%b %d %Y %I:%M:%S:%f%p")
                            date_str = dt.strftime('%Y-%m-%d')
                            time_str = dt.strftime('%H:%M')
                            curs.execute(
                                """INSERT INTO performance (
                                    performance_id, show_id, building_id, hallname,
                                    date, time, minprice, maxprice, freeplaces,
                                    building_name, hall_id
                                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                                ON CONFLICT (performance_id) DO NOTHING;""",
                                (
                                    item['IdPerformance'], item['IdShow'], item['IdBuilding'],
                                    item['HallName'], date_str, time_str,
                                    item['MinPrice'], item['MaxPrice'], item['FreePlace'],
                                    item['BuildingName'], item['IdHall'],
                                ),
                            )
                        except Exception as e:
                            logger.exception(f"Ошибка при вставке performance: {e}")
        except Exception as e:
            logger.exception("Ошибка all_performances_request")
            asyncio.run_coroutine_threadsafe(
                notify_admin(f'Ошибка all_performances_request\n{e}'), loop)
    else:
        asyncio.run_coroutine_threadsafe(
            notify_admin(f'Ошибка all_performances_request: код {response.status_code}'), loop)


# ─── Регистрация / ФИО пользователя ──────────────────────────────────────────
def user_reg(user_id: int, chat_id: int) -> dict:
    with psycopg2.connect(db_path) as conn:
        with conn.cursor() as curs:
            curs.execute(
                "SELECT user_id, max_chat_id, buyer_id, name, surname, patronymic, agreement "
                "FROM users WHERE user_id = %s and max_chat_id = %s;",
                (user_id, chat_id),
            )
            user = curs.fetchone()

            if user is None:
                params = {
                    "sp": "Wga_Autorize",
                    "Login": f'a{user_id}c',
                    "Name": f"maxid{user_id}",
                    "IdDocument": "3165",
                    "df": "J",
                }
                # buyer_id = 998277
                try:
                    response = requests.get(url_kino_baza, params=params)
                    response_data = response.json()
                    buyer_id = response_data['IdClient']
                except Exception:
                    a = decode_unicode(response.text)
                    logger.exception(f"Wga_Autorize: {a}")
                    buyer_id = 998277

                curs.execute(
                    "INSERT INTO users (user_id, max_chat_id, buyer_id) VALUES (%s, %s, %s) "
                    "RETURNING user_id, max_chat_id, buyer_id, name, surname, patronymic, agreement;",
                    (user_id, chat_id, buyer_id),
                )
                user = curs.fetchone()

            return {
                "user_id": user[0],
                "max_chat_id": user[1],
                "buyer_id": user[2],
                "name": user[3],
                "surname": user[4],
                "patronymic": user[5],
                "agreement": user[6],
            }


def user_fio_save(user_id: int, chat_id: int, name: str, surname: str, patronymic: str, agreement: bool) -> dict:
    with psycopg2.connect(db_path) as conn:
        with conn.cursor() as curs:
            curs.execute("SELECT user_id, max_chat_id FROM users WHERE user_id = %s and max_chat_id = %s;",
                         (user_id, chat_id,))
            if curs.fetchone() is None:
                params = {
                    "sp": "Wga_Autorize",
                    "Login": f'a{user_id}c',
                    "Name": f"maxid{user_id}",
                    "IdDocument": "3165",
                    "df": "J",
                }
                try:
                    buyer_id = requests.get(url_kino_baza, params=params).json()['IdClient']
                except Exception:
                    logger.exception("user_fio_save")
                    buyer_id = 998277
                curs.execute(
                    "INSERT INTO users (user_id, max_chat_id, buyer_id) VALUES (%s, %s, %s);",
                    (user_id, chat_id, buyer_id),
                )

            curs.execute(
                """UPDATE users SET name=%s, surname=%s, patronymic=%s, agreement=%s
                   WHERE user_id=%s and max_chat_id = %s
                   RETURNING user_id, max_chat_id, buyer_id, name, surname, patronymic, agreement;""",
                (name, surname, patronymic, agreement, user_id, chat_id),
            )
            user = curs.fetchone()
            conn.commit()
            return {
                "user_id": user[0],
                "max_chat_id": user[1],
                "buyer_id": user[2],
                "name": user[3],
                "surname": user[4],
                "patronymic": user[5],
                "agreement": user[6],
            }


# ─── Валидация Почта Банк ─────────────────────────────────────────────────────
def validate_pochta_bank(buyer_info, rrn, event_id, place_id, event_session_timestamp,
                         term_inst_id="11132", organization_id="31698",
                         client_buy_ip_address="0.0.0.0"):
    buyer_hash = hashlib.sha512(buyer_info.encode('utf-8')).hexdigest()
    json_data = {
        "buyer": buyer_hash,
        "termInstId": term_inst_id,
        "rrn": rrn,
        "eventId": event_id,
        "placeId": place_id,
        "organizationId": organization_id,
        "eventSessionTimestamp": event_session_timestamp,
        "clientBuyIpAddress": client_buy_ip_address,
    }
    headers = {
        "Authorization": pochta_bank_token,
        'User-Agent': 'Mozilla/5.0',
    }
    response = requests.post(
        validation_url, headers=headers,
        params={"online": "false"}, json=json_data
    )
    if response.status_code == 200:
        resp = response.json()
        result = resp.get("result")
        if result == "ACCEPTED":
            return True, ""
        if result == "DECLINED":
            reason_map = {
                "PERSONAL_DATA": "Несоответствие ФИО покупателя ФИО держателя карты",
                "EVENT": "Проблема с параметрами мероприятия",
                "MULTISESSION": "Попытка купить более одного билета на один сеанс",
                "PRICE": "Некорректная стоимость билета",
                "OTHER": "",
            }
            return False, reason_map.get(resp.get("reason", "OTHER"), "")
    else:
        logger.info(f"validate_pochta_bank {response.status_code=}, {response.text=}")
        return False, "400"


# ─── Отправка XML в Министерство ─────────────────────────────────────────────
async def send_xml_to_ekinobilet(performance_data, show_data, payment, fond_kino_id,
                                 place, row, price, payment_id, order_id):
    try:
        import xml.etree.ElementTree as ET
        await notify_admin(f'Проверка отправки')

        building_name = performance_data[6]
        id_procult = show_data[2]
        seans_date = f"{performance_data[8].replace('-', '')} {performance_data[9]}"
        pu_number = show_data[0]
        film_name = show_data[1]
        hall_name = performance_data[3]

        delta = datetime.timedelta(hours=7)
        today = datetime.datetime.now(datetime.timezone.utc) + delta
        sale_date = today.strftime('%Y%m%d %H:%M:%S')
        doc_date = today.strftime('%Y%m%d_%H%M%S')
        rrn = payment.authorization_details.rrn

        root_el = ET.Element('seans')
        root_el.set('ver', '3.2.0')
        root_el.set('org_id', str(fond_kino_id))
        root_el.set('showroom', str(hall_name))
        root_el.set('seans_date', str(seans_date))
        root_el.set('pu_number', str(pu_number))
        root_el.set('format', '2D')
        root_el.set('seans_title', str(film_name))
        root_el.set('event_id', str(id_procult))

        form = ET.SubElement(root_el, 'form')
        form.set('place_x', str(place))
        form.set('place_y', str(row))
        form.set('section', str(hall_name))
        form.set('price', str(price))
        form.set('discount', '0')
        form.set('ticket_type', 'Основной')
        form.set('sale_date', str(sale_date))
        form.set('subscription', 'false')
        form.set('online', 'true')
        form.set('webtax_included', 'false')
        form.set('payment_type', '1')
        form.set('payment_id', str(rrn))
        form.set('terminal_id', '26485891')
        form.set('terminal_owner', '5402052576')

        xml_data = ET.tostring(root_el, encoding='utf-8')
        xml_dir = os.path.join(root_path, "xml_files")
        os.makedirs(xml_dir, exist_ok=True)
        xml_file = os.path.join(xml_dir, f'ekb_{fond_kino_id}_{doc_date}145.xml')
        with open(xml_file, 'wb') as f:
            f.write(xml_data)

        try:
            await notify_admin(f'xml_file_name {order_id}\n{xml_file}')
            await send_document_sync(ADMIN_ID_CHAT, ADMIN_ID, xml_file, caption=f"XML файл заказа {order_id}",
                                     filename=os.path.basename(xml_file))
        except Exception:
            logger.exception("send_document_sync")

        auth = {'login': '505@mirkino.pro', 'password': 'pukugk'}
        with open(xml_file, 'rb') as f:
            for _ in range(5):
                try:
                    response = requests.post(
                        'https://ekinobilet.ru/ekbs/upload.aspx',
                        data=auth,
                        files={'XMLfile': f}
                    )
                    break
                except Exception:
                    time.sleep(7)

        text_resp = response.content.decode('utf-8')
        if 'error' not in text_resp:
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute(
                        "UPDATE orders SET report_sented = 1 WHERE payment_id = %s", (payment_id,)
                    )
        else:
            await notify_admin(
                f'!!!!Ошибка отправки в министерство\n'
                f'order_id {order_id} ошибка {text_resp} файл {xml_file}'
            )
    except Exception as e:
        logger.exception("Ошибка send_xml_to_ekinobilet")
        await notify_admin(f'!!!!Ошибка отправки XML\n{e} order_id {order_id}')


# def _inline_url(rows: list[list[tuple[str, str]]]):
#     return ButtonsPayload(
#         buttons=[
#             [
#                 LinkButton(
#                     text=text,
#                     url=url,
#                 )
#                 for text, url in row
#             ]
#             for row in rows
#         ]
#     ).pack()


# ─── Проверка статуса платежа ─────────────────────────────────────────────────
async def check_payment_status(payment_id: str, report: bool = True):
    is_succeeded = False

    Configuration.account_id = int(youkassa_shop_id)
    Configuration.secret_key = youkassa_secret_key

    payment = Payment.find_one(payment_id)
    payment_status = payment.status

    with psycopg2.connect(db_path) as conn:
        with conn.cursor() as curs:
            curs.execute("SELECT * FROM orders WHERE payment_id = %s", (payment_id,))
            order_data = curs.fetchone()
            curs.execute(
                "SELECT name, surname, patronymic, max_chat_id FROM users WHERE user_id = %s;",
                (order_data[1],)
            )
            user = curs.fetchone()

    fio = f"{user[1]} {user[0]} {user[2]}".strip()
    order_id = order_data[0]
    place_id = order_data[4]
    price = order_data[5]
    user_id = order_data[1]
    chat_id = user[3]
    row = order_data[11]
    place = order_data[12]
    performance_id = order_data[3]
    payment_msg_id = order_data[15] if len(order_data) > 15 else None

    if payment_status not in ["canceled", "succeeded", "pending", "waiting_for_capture"]:
        await notify_admin(
            f'!!!!!Ошибка проверки статуса заказа\n'
            f'{payment.status} {payment.id} {order_id}'
        )
        return is_succeeded

    if payment_status == "waiting_for_capture":
        Payment.capture(payment_id)

    elif payment_status == 'canceled':
        params = {"sp": "WgA_SetOrderToNull", "idOrder": order_id, "df": "J"}
        for _ in range(5):
            try:
                requests.get(url_kino_baza, params=params)
                break
            except Exception as e:
                await notify_admin(f'!!!!Ошибка отмены заказа {order_id}: {e}')
                time.sleep(7)

        with psycopg2.connect(db_path) as conn:
            with conn.cursor() as curs:
                curs.execute("UPDATE orders SET status = 0 WHERE payment_id = %s", (payment_id,))

        await send_message_sync(
            chat_id,
            user_id,
            f'Заказ не был оплачен вовремя.',
        )
        return "canceled"

    elif payment_status == 'succeeded':
        is_succeeded = True
        is_fk_report_send = True

        try:
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute(
                        "SELECT payment_id, payment_link, user_id, price, row, place, performance_id "
                        "FROM orders WHERE order_id = %s", (order_id,)
                    )
                    order = curs.fetchone()

                    curs.execute(
                        "SELECT hallname, date, time, show_id, building_id "
                        "FROM performance WHERE performance_id = %s", (order[6],)
                    )
                    performance = curs.fetchone()

                    curs.execute("SELECT name FROM show WHERE show_id = %s", (performance[3],))
                    show = curs.fetchone()

                    curs.execute("SELECT city FROM cinemas WHERE building_id = %s", (performance[4],))
                    cinema = curs.fetchone()

                    user_id = order[2]
                    price = order[3]
                    row = order[4]
                    place = order[5]
                    hallname = performance[0]
                    date = performance[1]
                    time_ = performance[2]
                    name = show[0]
                    city = cinema[0]

                    months = {
                        1: "января", 2: "февраля", 3: "марта", 4: "апреля",
                        5: "мая", 6: "июня", 7: "июля", 8: "августа",
                        9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
                    }
                    date_obj = datetime.datetime.strptime(date, "%Y-%m-%d")
                    date_ru = f"{date_obj.day} {months[date_obj.month]} {date_obj.year}"
                    date_time_ru = f"{date_ru} {time_}"

                    # Генерация PDF-билета
                    template_path = os.path.join(root_path, "Шаблон билета.docx")
                    doc = Document(template_path)

                    now_plus7 = datetime.datetime.utcnow() + datetime.timedelta(hours=7)
                    purchase_date_ru = (
                        f"{now_plus7.day} {months[now_plus7.month]} {now_plus7.year} "
                        f"{now_plus7.hour:02}:{now_plus7.minute:02}"
                    )

                    replace_dict = {
                        "{city}": city, "{cinema}": hallname,
                        "{number}": str(order_id), "{date}": purchase_date_ru,
                        "{cinema_date}": date_time_ru, "{name}": name,
                        "{row}": str(row), "{place}": str(place),
                        "{price}": str(price), "{fio}": fio,
                    }

                    def replace_placeholder_safe(paragraph, rd):
                        full_text = ''.join(r.text for r in paragraph.runs if r.text)
                        replaced = False
                        for key, value in rd.items():
                            if key in full_text:
                                full_text = full_text.replace(key, value)
                                replaced = True
                        if not replaced:
                            return
                        idx = 0
                        for run in paragraph.runs:
                            if run.text:
                                length = len(run.text)
                                run.text = full_text[idx:idx + length]
                                idx += length
                        if idx < len(full_text):
                            for run in reversed(paragraph.runs):
                                if run.text:
                                    run.text += full_text[idx:]
                                    break

                    for paragraph in doc.paragraphs:
                        replace_placeholder_safe(paragraph, replace_dict)
                    for table in doc.tables:
                        for trow in table.rows:
                            for cell in trow.cells:
                                for paragraph in cell.paragraphs:
                                    replace_placeholder_safe(paragraph, replace_dict)

                    docx_path = f"/tmp/Билет_№{order_id}.docx"
                    pdf_path = f"/tmp/Билет_№{order_id}.pdf"
                    doc.save(docx_path)

                    subprocess.run(
                        ["libreoffice", "--headless", "--convert-to", "pdf",
                         docx_path, "--outdir", "/tmp"],
                        check=True,
                    )

                    msg_text = (
                        "✅ Заказ оплачен\n\n"
                        f"Заказ №: {order_id}\n"
                        f"Сеанс: {name}\n"
                        f"Кинотеатр: {hallname} ({city})\n"
                        f"Дата и время сеанса: {date_time_ru}\n"
                        f"Ряд / Место: {row} / {place}\n"
                        f"Цена: {price} р.\n\n"
                        "👇 Ваш билет ниже 👇"
                    )
                    await send_message_sync(chat_id, user_id, msg_text)
                    await send_document_sync(
                        chat_id, user_id, pdf_path,
                        caption=f"Ваш билет №{order_id}",
                        filename=f"Билет_№{order_id}.pdf",
                    )

                    try:
                        os.remove(docx_path)
                        os.remove(pdf_path)
                    except Exception:
                        logger.exception("Ошибка удаления временных файлов")

                    # Добавляем оплату в МирКино
                    params = {
                        "sp": "WgA_AddPayment",
                        "IdOrder": order_id,
                        "Amount": price,
                        "IdPaymentMethod": 11,
                        "idUser": 1,
                        "df": "J",
                    }
                    for _ in range(5):
                        try:
                            resp = requests.get(url_kino_baza, params=params)
                            break
                        except Exception:
                            time.sleep(7)
                    payment_kino = resp.json()
                    kino_add_payment_id = payment_kino['IdPayment']

                    with psycopg2.connect(db_path) as conn2:
                        with conn2.cursor() as curs2:
                            try:
                                curs2.execute(
                                    "UPDATE orders SET status = 1, kino_add_payment_id = %s "
                                    "WHERE payment_id = %s",
                                    (int(kino_add_payment_id), payment_id),
                                )
                            except (ValueError, TypeError):
                                curs2.execute(
                                    "UPDATE orders SET status = 1 WHERE payment_id = %s",
                                    (payment_id,),
                                )

                            curs2.execute(
                                "SELECT * FROM performance WHERE performance_id = %s",
                                (performance_id,)
                            )
                            performance_data = curs2.fetchone()

                            curs2.execute(
                                "SELECT fond_kino_id FROM cinemas WHERE building_id = %s",
                                (performance_data[2],)
                            )
                            fond_kino_id = curs2.fetchone()[0]

                            curs2.execute(
                                "SELECT pu_number, name, id_procult FROM show WHERE show_id = %s",
                                (performance_data[1],)
                            )
                            show_data = curs2.fetchone()

                            curs2.execute(
                                "SELECT report_sented FROM orders WHERE payment_id = %s", (payment_id,)
                            )
                            is_fk_report_send = curs2.fetchone()[0]

        except TypeError:
            err_msg = (
                f'!!!!Ошибка оформления заказа\n'
                f'order_id {order_data[0]}\nuser_id {order_data[1]}\n'
                f'performance {order_data[3]}\nplace_id {order_data[4]}\n'
                f'ряд {order_data[11]}\nместо {order_data[12]}\n'
                f'payment_id {order_data[6]}'
            )
            await notify_admin(err_msg)
            if report:
                await notify_admin(err_msg, admin_2=True)
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute(
                        "UPDATE orders SET status = 4 WHERE payment_id = %s", (payment_id,)
                    )
        except Exception as e:
            logger.exception("Ошибка при обработке succeeded")
            await notify_admin(f'!!!!Ошибка оформления succeeded\n{e}')

        if not is_fk_report_send:
            await send_xml_to_ekinobilet(
                performance_data, show_data, payment, fond_kino_id,
                place, row, price, payment_id, order_id,
            )

        # Валидация Почта Банк
        try:
            rrn = payment.authorization_details.rrn
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute(
                        "SELECT date, time, show_id FROM performance WHERE performance_id = %s",
                        (performance_id,)
                    )
                    perf = curs.fetchone()
                    curs.execute("SELECT id_procult FROM show WHERE show_id = %s", (perf[2],))
                    event_id = curs.fetchone()[0]

            date_time_obj = datetime.datetime.strptime(f'{perf[0]} {perf[1]}', '%Y-%m-%d %H:%M')
            event_session_timestamp = int(time.mktime(date_time_obj.timetuple()))

            ok, text = validate_pochta_bank(
                buyer_info=fio, rrn=rrn, event_id=event_id,
                place_id=place_id, event_session_timestamp=event_session_timestamp,
            )
            if not ok:
                if text == "400":
                    return is_succeeded
                await send_xml_to_ekinobilet(
                    performance_data, show_data, payment, fond_kino_id,
                    place, row, -price, payment_id, order_id,
                )
                await send_message_sync(
                    chat_id,
                    user_id,
                    f"Извините, возникла ошибка, деньги вернутся на вашу карту.\n{text}",
                )
        except Exception:
            logger.exception("Ошибка валидации Почта Банк")

    return is_succeeded


# ─── Разблокировка мест ───────────────────────────────────────────────────────
def unblock_all(user_id: int, performance_id: int, place_id, loop):
    with psycopg2.connect(db_path) as conn:
        with conn.cursor() as curs:
            if place_id == 'all':
                curs.execute(
                    "SELECT performance_id, place_id, buyer_id, order_id FROM orders "
                    "WHERE user_id = %s AND status = 2;",
                    (user_id,)
                )
            else:
                curs.execute(
                    "SELECT performance_id, place_id, buyer_id, order_id FROM orders "
                    "WHERE user_id = %s AND status = 2 AND place_id != %s;",
                    (user_id, place_id)
                )
            orders_to_close = curs.fetchall()

            for order in orders_to_close:
                params = {"sp": "WgA_SetOrderToNull", "idOrder": order[3], "df": "J"}
                for _ in range(5):
                    try:
                        requests.get(url_kino_baza, params=params)
                        break
                    except Exception as e:
                        asyncio.run_coroutine_threadsafe(notify_admin(f'!!!!Ошибка unblock_all: {e}'), loop)
                        time.sleep(7)
                curs.execute(
                    "UPDATE orders SET status = 0 "
                    "WHERE performance_id = %s AND place_id = %s AND buyer_id = %s AND user_id = %s",
                    (order[0], order[1], order[2], user_id)
                )


def unblock_5_min(loop):
    with psycopg2.connect(db_path) as conn:
        with conn.cursor() as curs:
            curs.execute(
                "SELECT performance_id, place_id, buyer_id, payment_id, order_id, user_id "
                "FROM orders WHERE status = 2 AND place_locked_time < %s;",
                (time.time() - 300,)
            )
            to_unblock = curs.fetchall()

    for order in to_unblock:
        if order[4] is None:
            with psycopg2.connect(db_path) as conn:
                with conn.cursor() as curs:
                    curs.execute(
                        "UPDATE orders SET status = 0 WHERE user_id = %s AND performance_id = %s",
                        (order[5], order[0])
                    )
            continue

        logger.info(order)
        params = {"sp": "WgA_SetOrderToNull", "idOrder": order[4], "df": "J"}
        for _ in range(5):
            try:
                response = requests.get(url_kino_baza, params=params)
                logger.info(decode_unicode(response.text))
                break
            except Exception as e:
                asyncio.run_coroutine_threadsafe(notify_admin(f'!!!!Ошибка unblock_5_min: {e}'), loop)
                time.sleep(7)

        with psycopg2.connect(db_path) as conn:
            with conn.cursor() as curs:
                curs.execute("UPDATE orders SET status = 0 WHERE order_id = %s", (order[4],))


def decode_unicode(data):
    if isinstance(data, str):
        return data.encode('utf-8').decode('unicode_escape')
    elif isinstance(data, dict):
        return {k: decode_unicode(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [decode_unicode(v) for v in data]
    return data
